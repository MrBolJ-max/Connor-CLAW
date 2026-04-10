"""
SYN Systems — Document Ingestion & Knowledge Base
Step 01 of the SYN Systems onboarding process.

Supports: PDF, DOCX, TXT, Markdown, raw text paste
Extracts: SOPs, product docs, past conversations, brand voice
Stores: In local vector-ready JSON + optional Pinecone index
"""

import json
import os
import re
import hashlib
from datetime import datetime
from pathlib import Path
from dataclasses import dataclass, field, asdict
from typing import Optional
from loguru import logger
from core.claude_client import claude
from core.config import settings

DATA_DIR = Path(__file__).parent.parent / "data" / "knowledge_bases"
DATA_DIR.mkdir(parents=True, exist_ok=True)


@dataclass
class Document:
    doc_id: str
    client_id: str
    doc_type: str           # sop, product, conversation, brand_voice, faq, contract
    title: str
    raw_text: str
    summary: str = ""
    key_facts: list[str] = field(default_factory=list)
    brand_voice_notes: str = ""
    escalation_triggers: list[str] = field(default_factory=list)
    chunk_count: int = 0
    ingested_at: str = field(default_factory=lambda: datetime.utcnow().isoformat())
    file_hash: str = ""

    def to_dict(self) -> dict:
        return asdict(self)


@dataclass
class ClientKnowledgeBase:
    client_id: str
    client_name: str
    industry: str
    documents: list[Document] = field(default_factory=list)
    brand_voice: str = ""
    key_processes: list[str] = field(default_factory=list)
    escalation_rules: list[str] = field(default_factory=list)
    agent_instructions: str = ""
    created_at: str = field(default_factory=lambda: datetime.utcnow().isoformat())
    last_updated: str = field(default_factory=lambda: datetime.utcnow().isoformat())

    @property
    def doc_count(self) -> int:
        return len(self.documents)

    def to_dict(self) -> dict:
        d = asdict(self)
        d["doc_count"] = self.doc_count
        return d


# ── Text Extraction ───────────────────────────────────────────────
def extract_text_from_pdf(file_path: str) -> str:
    try:
        import PyPDF2
        text_parts = []
        with open(file_path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                text_parts.append(page.extract_text() or "")
        return "\n".join(text_parts)
    except ImportError:
        logger.warning("PyPDF2 not installed — cannot parse PDF")
        return ""
    except Exception as e:
        logger.error(f"PDF extraction failed: {e}")
        return ""


def extract_text_from_docx(file_path: str) -> str:
    try:
        from docx import Document as DocxDocument
        doc = DocxDocument(file_path)
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    except ImportError:
        logger.warning("python-docx not installed — cannot parse DOCX")
        return ""
    except Exception as e:
        logger.error(f"DOCX extraction failed: {e}")
        return ""


def extract_text(source: str, file_path: Optional[str] = None) -> str:
    """Extract plain text from various sources."""
    if file_path:
        ext = Path(file_path).suffix.lower()
        if ext == ".pdf":
            return extract_text_from_pdf(file_path)
        elif ext in (".docx", ".doc"):
            return extract_text_from_docx(file_path)
        elif ext in (".txt", ".md", ".markdown"):
            return Path(file_path).read_text(encoding="utf-8")
    return source  # raw text passed directly


def chunk_text(text: str, chunk_size: int = 800, overlap: int = 100) -> list[str]:
    """Split text into overlapping chunks for embedding."""
    words = text.split()
    chunks = []
    start = 0
    while start < len(words):
        end = start + chunk_size
        chunk = " ".join(words[start:end])
        chunks.append(chunk)
        start += chunk_size - overlap
    return chunks


def file_hash(text: str) -> str:
    return hashlib.md5(text.encode()).hexdigest()[:12]


# ── Analysis Engine ───────────────────────────────────────────────
class DocumentAnalyser:
    """Uses Claude to extract structured insights from raw documents."""

    ANALYSIS_PROMPT = """Analyse this business document and extract structured information.

Return valid JSON:
{
  "summary": "<2-3 sentence summary>",
  "key_facts": ["<fact1>", "<fact2>", "<fact3>"],
  "brand_voice_notes": "<how they communicate: tone, style, formality>",
  "escalation_triggers": ["<situation requiring human>", "<another trigger>"],
  "processes": ["<key process 1>", "<key process 2>"],
  "faq_pairs": [{"q": "<question>", "a": "<answer>"}]
}"""

    def analyse(self, text: str, doc_type: str) -> dict:
        try:
            raw = claude.extract_json(
                system_prompt=self.ANALYSIS_PROMPT,
                user_message=f"Document type: {doc_type}\n\n---\n{text[:6000]}",
                max_tokens=1000,
            )
            return json.loads(raw)
        except (json.JSONDecodeError, Exception) as e:
            logger.error(f"Document analysis failed: {e}")
            return {
                "summary": text[:200],
                "key_facts": [],
                "brand_voice_notes": "",
                "escalation_triggers": [],
                "processes": [],
                "faq_pairs": [],
            }

    def synthesise_brand_voice(self, documents: list[Document]) -> str:
        """Synthesise a brand voice guide from all ingested documents."""
        notes = [d.brand_voice_notes for d in documents if d.brand_voice_notes]
        if not notes:
            return "Professional, helpful, and solution-focused."
        combined = "\n".join(notes)
        return claude.chat(
            system_prompt="You are a brand voice strategist. Synthesise a concise brand voice guide.",
            user_message=f"Synthesise these brand voice observations:\n{combined}",
            max_tokens=300,
        )

    def generate_agent_instructions(self, kb: "ClientKnowledgeBase") -> str:
        """Generate the system prompt for this client's AI agents."""
        return claude.chat(
            system_prompt=(
                "You are a prompt engineer specialising in AI agent configuration. "
                "Write clear, comprehensive system prompts for AI agents."
            ),
            user_message=(
                f"Generate a system prompt for an AI agent for:\n"
                f"Client: {kb.client_name}\n"
                f"Industry: {kb.industry}\n"
                f"Brand Voice: {kb.brand_voice}\n"
                f"Key Processes: {json.dumps(kb.key_processes)}\n"
                f"Escalation Rules: {json.dumps(kb.escalation_rules)}\n"
                f"Number of docs ingested: {kb.doc_count}\n\n"
                "The agent should sound like their best team member. Include: role, tone, "
                "key knowledge areas, what to escalate, and how to handle unknowns."
            ),
            max_tokens=600,
        )


# ── Knowledge Base Manager ────────────────────────────────────────
class KnowledgeBaseManager:
    """
    Manages the full lifecycle of client knowledge bases.
    Handles ingestion, storage, retrieval, and agent instruction generation.
    """

    def __init__(self):
        self.analyser = DocumentAnalyser()
        self._cache: dict[str, ClientKnowledgeBase] = {}

    def _kb_path(self, client_id: str) -> Path:
        return DATA_DIR / f"{client_id}.json"

    def save(self, kb: ClientKnowledgeBase):
        kb.last_updated = datetime.utcnow().isoformat()
        path = self._kb_path(kb.client_id)
        path.write_text(json.dumps(kb.to_dict(), indent=2), encoding="utf-8")
        self._cache[kb.client_id] = kb
        logger.info(f"KB saved: {kb.client_name} ({kb.doc_count} docs)")

    def load(self, client_id: str) -> Optional[ClientKnowledgeBase]:
        if client_id in self._cache:
            return self._cache[client_id]
        path = self._kb_path(client_id)
        if not path.exists():
            return None
        data = json.loads(path.read_text(encoding="utf-8"))
        docs = [Document(**d) for d in data.pop("documents", [])]
        data.pop("doc_count", None)
        kb = ClientKnowledgeBase(**data)
        kb.documents = docs
        self._cache[client_id] = kb
        return kb

    def list_clients(self) -> list[dict]:
        clients = []
        for path in DATA_DIR.glob("*.json"):
            try:
                data = json.loads(path.read_text(encoding="utf-8"))
                clients.append({
                    "client_id": data.get("client_id"),
                    "client_name": data.get("client_name"),
                    "industry": data.get("industry"),
                    "doc_count": data.get("doc_count", 0),
                    "last_updated": data.get("last_updated"),
                })
            except Exception:
                continue
        return sorted(clients, key=lambda x: x.get("last_updated", ""), reverse=True)

    def create_client(self, client_id: str, client_name: str, industry: str) -> ClientKnowledgeBase:
        kb = ClientKnowledgeBase(
            client_id=client_id,
            client_name=client_name,
            industry=industry,
        )
        self.save(kb)
        logger.info(f"New KB created: {client_name}")
        return kb

    def ingest_document(
        self,
        client_id: str,
        doc_type: str,
        title: str,
        text: str = "",
        file_path: Optional[str] = None,
    ) -> Document:
        """
        Ingest a document into a client's knowledge base.
        Step 01 of SYN Systems onboarding.
        """
        kb = self.load(client_id)
        if kb is None:
            raise ValueError(f"No knowledge base found for client '{client_id}'")

        raw_text = extract_text(text, file_path)
        if not raw_text.strip():
            raise ValueError("Document appears to be empty after extraction")

        fhash = file_hash(raw_text)
        # Dedup check
        for existing in kb.documents:
            if existing.file_hash == fhash:
                logger.info(f"Duplicate document skipped: {title}")
                return existing

        logger.info(f"Analysing document: {title} ({doc_type})")
        analysis = self.analyser.analyse(raw_text, doc_type)
        chunks = chunk_text(raw_text)

        doc = Document(
            doc_id=f"{client_id}_{fhash}",
            client_id=client_id,
            doc_type=doc_type,
            title=title,
            raw_text=raw_text[:5000],  # store first 5k chars
            summary=analysis.get("summary", ""),
            key_facts=analysis.get("key_facts", []),
            brand_voice_notes=analysis.get("brand_voice_notes", ""),
            escalation_triggers=analysis.get("escalation_triggers", []),
            chunk_count=len(chunks),
            file_hash=fhash,
        )

        kb.documents.append(doc)

        # Update KB-level aggregates
        all_processes = analysis.get("processes", [])
        kb.key_processes = list(set(kb.key_processes + all_processes))[:20]
        all_triggers = analysis.get("escalation_triggers", [])
        kb.escalation_rules = list(set(kb.escalation_rules + all_triggers))[:15]

        # Re-synthesise brand voice with new data
        kb.brand_voice = self.analyser.synthesise_brand_voice(kb.documents)

        # Regenerate agent instructions
        kb.agent_instructions = self.analyser.generate_agent_instructions(kb)

        self.save(kb)
        logger.info(f"Document ingested: {title} | {len(chunks)} chunks | KB: {kb.client_name}")
        return doc

    def search(self, client_id: str, query: str, top_k: int = 5) -> list[str]:
        """Simple keyword search across a client's KB (Pinecone optional for semantic)."""
        kb = self.load(client_id)
        if kb is None:
            return []
        query_lower = query.lower()
        results = []
        for doc in kb.documents:
            score = sum(1 for word in query_lower.split() if word in doc.raw_text.lower())
            if score > 0:
                results.append((score, doc.summary, doc.title))
        results.sort(reverse=True)
        return [f"[{title}] {summary}" for _, summary, title in results[:top_k]]

    def get_agent_system_prompt(self, client_id: str) -> str:
        """Retrieve the generated system prompt for a client's agent."""
        kb = self.load(client_id)
        if kb is None:
            return "You are a helpful AI assistant."
        return kb.agent_instructions or "You are a helpful AI assistant."


kb_manager = KnowledgeBaseManager()
